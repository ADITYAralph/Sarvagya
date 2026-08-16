"""
Sarvagya Document Service — DOCX & PDF Resume Extraction
=========================================================
Extracts clean, structured text from .docx and .pdf files.
Returns the same interface for both formats.
Performs content validation to confirm the document is an actual resume.
"""

import io
import re
import logging
from typing import Dict, Any

logger = logging.getLogger("sarvagya.doc_service")

# Standard section headers expected in a genuine resume
STANDARD_RESUME_HEADERS = [
    "experience", "work", "projects", "education", "skills",
    "summary", "employment", "qualifications", "certifications",
    "objective", "profile", "career", "achievements", "internship",
    "training", "courses", "languages", "hobbies", "awards",
]

# Accepted file extensions
ALLOWED_EXTENSIONS = {".docx", ".pdf"}


def validate_and_extract(file_bytes: bytes, filename: str = "") -> Dict[str, Any]:
    """
    Unified entry point. Routes to DOCX or PDF extractor based on filename extension.
    Returns a dict with: is_valid, text, word_count, char_count, line_count, error_reason.
    """
    ext = ("."+filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""
    if ext == ".docx":
        return validate_and_extract_docx(file_bytes, filename)
    elif ext == ".pdf":
        return validate_and_extract_pdf(file_bytes, filename)
    else:
        return _error_result(
            f"Invalid file type '{ext}'. Only .docx (Word) and .pdf files are accepted. "
            "Please upload your resume in one of these formats."
        )


def validate_and_extract_docx(file_bytes: bytes, filename: str = "") -> Dict[str, Any]:
    """
    Extracts text from a .docx file and validates it as a genuine resume.
    """
    try:
        from docx import Document  # python-docx
    except ImportError:
        return _error_result("python-docx is not installed. Run: pip install python-docx")


    # ── Parse DOCX ─────────────────────────────────────────────
    try:
        doc_io = io.BytesIO(file_bytes)
        doc = Document(doc_io)
    except Exception as e:
        logger.error(f"Failed to open DOCX: {e}")
        return _error_result(
            "File could not be opened as a Word document (.docx). "
            "Ensure the file is a valid, non-corrupted .docx resume."
        )

    # ── Extract text preserving structure ──────────────────────
    lines = []

    # Paragraphs (main body)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Tables (skill tables, two-column layouts, etc.)
    for table in doc.tables:
        for row in table.rows:
            row_text = "  |  ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                lines.append(row_text)

    extracted_text = "\n".join(lines).strip()

    # ── Basic extraction metrics ───────────────────────────────
    words = re.findall(r"\b[A-Za-z0-9+#.\-]+\b", extracted_text)
    word_count = len(words)
    char_count = len(extracted_text)
    line_count = len(lines)

    logger.info(
        f"[DOCX] File: {filename} | Words: {word_count} | "
        f"Chars: {char_count} | Lines: {line_count}"
    )
    logger.debug(f"[DOCX] Text preview:\n{extracted_text[:400]}\n...")

    # ── Guard: too short ───────────────────────────────────────
    if word_count < 50:
        return {
            "is_valid": False,
            "text": extracted_text,
            "word_count": word_count,
            "char_count": char_count,
            "line_count": line_count,
            "error_reason": (
                f"Document contains only {word_count} words — too short to be a valid resume. "
                "A genuine resume should have at least 50 words of readable content."
            ),
        }

    # ── Guard: resume content check ────────────────────────────
    text_lower = extracted_text.lower()
    matched_headers = [h for h in STANDARD_RESUME_HEADERS if h in text_lower]

    if len(matched_headers) < 2:
        return {
            "is_valid": False,
            "text": extracted_text,
            "word_count": word_count,
            "char_count": char_count,
            "line_count": line_count,
            "error_reason": (
                "This document does not appear to be a resume. "
                "Standard resume sections (Experience, Education, Skills, Projects) were not detected. "
                "Please upload your actual resume in .docx or .pdf format."
            ),
        }

    # ── All checks passed ──────────────────────────────────────
    return {
        "is_valid": True,
        "text": extracted_text,
        "word_count": word_count,
        "char_count": char_count,
        "line_count": line_count,
        "error_reason": "",
    }


def validate_and_extract_pdf(file_bytes: bytes, filename: str = "") -> Dict[str, Any]:
    """
    Extracts text from a .pdf file and validates it as a genuine resume.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        return _error_result("pypdf is not installed. Run: pip install pypdf")

    try:
        pdf_io = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_io)
        pages_text = []
        for idx, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                clean_page = "\n".join(line.rstrip() for line in page_text.splitlines() if line.strip())
                pages_text.append(clean_page)
        extracted_text = "\n\n".join(pages_text).strip()
    except Exception as e:
        logger.error(f"Failed to parse PDF '{filename}': {e}")
        return _error_result(
            f"File could not be read as a PDF. Ensure it is a text-searchable (not scanned/image-only) PDF resume."
        )

    words = re.findall(r"\b[A-Za-z0-9+#.\-]+\b", extracted_text)
    word_count = len(words)
    char_count = len(extracted_text)
    line_count = len(extracted_text.splitlines())

    logger.info(f"[PDF] File: {filename} | Words: {word_count} | Chars: {char_count} | Lines: {line_count}")

    if word_count < 50:
        return {
            "is_valid": False,
            "text": extracted_text,
            "word_count": word_count,
            "char_count": char_count,
            "line_count": line_count,
            "error_reason": (
                f"Document contains only {word_count} words — too short to be a valid resume. "
                "Ensure the PDF is text-searchable (not a scanned image). "
                "A genuine resume should have at least 50 words."
            ),
        }

    text_lower = extracted_text.lower()
    matched_headers = [h for h in STANDARD_RESUME_HEADERS if h in text_lower]
    if len(matched_headers) < 2:
        return {
            "is_valid": False,
            "text": extracted_text,
            "word_count": word_count,
            "char_count": char_count,
            "line_count": line_count,
            "error_reason": (
                "This document does not appear to be a resume. "
                "Standard resume sections (Experience, Education, Skills, Projects) were not detected. "
                "Please upload your actual resume PDF."
            ),
        }

    return {
        "is_valid": True,
        "text": extracted_text,
        "word_count": word_count,
        "char_count": char_count,
        "line_count": line_count,
        "error_reason": "",
    }


def _error_result(reason: str) -> Dict[str, Any]:
    """Return a standardised error dict."""
    return {
        "is_valid": False,
        "text": "",
        "word_count": 0,
        "char_count": 0,
        "line_count": 0,
        "error_reason": reason,
    }
